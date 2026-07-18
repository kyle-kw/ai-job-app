from __future__ import annotations

import contextlib
import pathlib
import re
import uuid
from typing import Any

try:
    from worker_protocol import now
except ImportError:
    from sidecar.worker_protocol import now


def extract_docx_text(path: pathlib.Path) -> str:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = Document(str(path))
    lines: list[str] = []
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            value = Paragraph(child, document).text.strip()
            if value:
                lines.append(value)
        elif isinstance(child, CT_Tbl):
            table = Table(child, document)
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                value = " | ".join(dict.fromkeys(cell for cell in cells if cell))
                if value:
                    lines.append(value)
    for section in document.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                value = paragraph.text.strip()
                if value and value not in lines:
                    lines.append(value)
            for table in container.tables:
                for row in table.rows:
                    value = " | ".join(
                        cell.text.strip().replace("\n", " ")
                        for cell in row.cells
                        if cell.text.strip()
                    )
                    if value and value not in lines:
                        lines.append(value)
    return "\n".join(lines)


def render_pdf_page(path: pathlib.Path, page_index: int, output_dir: pathlib.Path) -> pathlib.Path:
    import pypdfium2 as pdfium

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"page-{page_index + 1}.png"
    with contextlib.closing(pdfium.PdfDocument(str(path))) as document:
        with contextlib.closing(document[page_index]) as page:
            with contextlib.closing(page.render(scale=2.25)) as bitmap:
                with contextlib.closing(bitmap.to_pil()) as image:
                    image.save(output_path, format="PNG", optimize=True)
    return output_path


def extract_text(path: pathlib.Path) -> tuple[str, dict[str, Any] | None, list[dict[str, Any]]]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        from pypdf import PdfReader

        pages: list[dict[str, Any]] = []
        page_texts: list[str] = []
        reader = PdfReader(str(path))
        image_dir = path.parent / f"{path.stem}-pages"
        for index, page in enumerate(reader.pages):
            page_text = (page.extract_text() or "").strip()
            page_texts.append(page_text)
            page_data: dict[str, Any] = {"pageNumber": index + 1, "text": page_text}
            has_visual_content = bool(list(page.images))
            if len(page_text) < 30 and has_visual_content:
                page_data["imagePath"] = str(render_pdf_page(path, index, image_dir).resolve())
            pages.append(page_data)
        return "\n\n".join(page_texts), None, pages
    if suffix == ".docx":
        text = extract_docx_text(path)
        if len(text.strip()) < 20:
            raise RuntimeError("DOCX 中没有足够的可读取文本。")
        return text, None, []
    if suffix in {".yaml", ".yml"}:
        import yaml

        data = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
        return path.read_text(encoding="utf-8"), data, []
    raise RuntimeError("仅支持 PDF、DOCX、YAML 和 YML。")


def first_value(data: dict[str, Any], *keys: str) -> str:
    for key in keys:
        value = data.get(key)
        if value is not None:
            return str(value)
    return ""


def split_skill_items(value: Any) -> list[str]:
    if isinstance(value, list):
        values = [str(item).strip() for item in value]
    else:
        values = [item.strip() for item in re.split(r"[,，、|]", str(value or ""))]
    return list(dict.fromkeys(item for item in values if item))


def entry_dates(entry: dict[str, Any]) -> tuple[str, str]:
    start = first_value(entry, "start_date", "startDate")
    end = first_value(entry, "end_date", "endDate")
    combined = first_value(entry, "date")
    if combined and not (start or end):
        start, end = normalize_date_pair("", combined)
    return normalize_date_pair(start, end)


DATE_RANGE_PATTERN = re.compile(
    r"^\s*(\d{4}(?:[./\-年]\d{1,2}(?:月)?)?)\s*(?:-|–|—|至|到)\s*(\d{4}(?:[./\-年]\d{1,2}(?:月)?)?|至今|现在|present)\s*$",
    re.IGNORECASE,
)


def normalize_date_pair(start: Any, end: Any) -> tuple[str, str]:
    start_value = str(start or "").strip().strip("-–— ")
    end_value = str(end or "").strip().strip("-–— ")
    candidate = end_value if not start_value else start_value if not end_value else ""
    match = DATE_RANGE_PATTERN.fullmatch(candidate) if candidate else None
    if match:
        return match.group(1).strip(), match.group(2).strip()
    return start_value, end_value


def render_date_fields(item: dict[str, Any]) -> tuple[str | None, str | None]:
    start, end = normalize_date_pair(item.get("startDate"), item.get("endDate"))
    if not start and end:
        return rendercv_date(end), None
    return rendercv_date(start), rendercv_date(end, end_date=True)


def rendercv_date(value: str, *, end_date: bool = False) -> str | None:
    normalized = value.strip()
    if not normalized:
        return None
    if end_date and normalized.lower() in {"至今", "现在", "present", "current"}:
        return "present"
    match = re.fullmatch(
        r"(\d{4})(?:[./年-](\d{1,2})(?:月)?)?(?:[./日-](\d{1,2})(?:日)?)?", normalized
    )
    if not match:
        return normalized
    return "-".join(
        part.zfill(2) if index else part
        for index, part in enumerate(match.groups(default=""))
        if part
    )


def rendercv_phone(value: Any) -> str | None:
    phone = str(value or "").strip()
    digits = re.sub(r"\D", "", phone)
    if len(digits) == 11 and digits.startswith("1"):
        return f"+86{digits}"
    return phone or None


def display_degree(item: dict[str, Any]) -> str:
    degree = str(item.get("degree") or "").strip()
    if degree == "其他":
        return str(item.get("degreeDetail") or "").strip() or degree
    return degree


def profile_from_yaml(data: dict[str, Any], file_name: str) -> dict[str, Any]:
    cv = data.get("cv") if isinstance(data.get("cv"), dict) else data
    sections = cv.get("sections") if isinstance(cv.get("sections"), dict) else {}
    experience_entries: list[dict[str, Any]] = []
    education_entries: list[dict[str, Any]] = []
    professional_skills: list[dict[str, Any]] = []
    projects: list[dict[str, Any]] = []
    certifications: list[dict[str, Any]] = []
    summary = ""
    for section_name, entries in sections.items():
        name = str(section_name).lower()
        entries = entries if isinstance(entries, list) else []
        if any(word in name for word in ["education", "教育"]):
            for entry in entries:
                if isinstance(entry, dict):
                    start_date, end_date = entry_dates(entry)
                    education_entries.append(
                        {
                            "institution": first_value(entry, "institution"),
                            "area": first_value(entry, "area"),
                            "degree": first_value(entry, "degree"),
                            "startDate": start_date,
                            "degreeDetail": first_value(entry, "degree_detail", "degreeDetail"),
                            "endDate": end_date,
                            "highlights": [str(item) for item in entry.get("highlights", [])],
                        }
                    )
        elif any(word in name for word in ["experience", "工作", "职业经历"]) and not any(
            word in name for word in ["project", "项目"]
        ):
            for entry in entries:
                if isinstance(entry, dict):
                    start_date, end_date = entry_dates(entry)
                    experience_entries.append(
                        {
                            "company": first_value(entry, "company", "institution"),
                            "position": first_value(entry, "position", "title"),
                            "location": first_value(entry, "location"),
                            "startDate": start_date,
                            "endDate": end_date,
                            "highlights": [str(item) for item in entry.get("highlights", [])],
                        }
                    )
        elif "skill" in name or "技能" in name:
            for entry in entries:
                if isinstance(entry, str):
                    professional_skills.append(
                        {
                            "id": str(uuid.uuid4()),
                            "label": "核心技能",
                            "items": split_skill_items(entry),
                        }
                    )
                elif isinstance(entry, dict):
                    professional_skills.append(
                        {
                            "id": str(uuid.uuid4()),
                            "label": first_value(entry, "label") or "专业技能",
                            "items": split_skill_items(entry.get("details") or entry.get("items")),
                        }
                    )
        elif any(word in name for word in ["project", "项目"]):
            for entry in entries:
                if isinstance(entry, dict):
                    start_date, end_date = entry_dates(entry)
                    projects.append(
                        {
                            "id": str(uuid.uuid4()),
                            "name": first_value(entry, "name", "title"),
                            "summary": first_value(entry, "summary"),
                            "startDate": start_date,
                            "endDate": end_date,
                            "highlights": [str(item) for item in entry.get("highlights", [])],
                        }
                    )
        elif any(word in name for word in ["certification", "certificate", "证书", "资质"]):
            for entry in entries:
                if isinstance(entry, dict):
                    certifications.append(
                        {
                            "id": str(uuid.uuid4()),
                            "name": first_value(entry, "name", "title"),
                            "issuer": first_value(entry, "issuer", "institution"),
                            "date": first_value(entry, "date"),
                        }
                    )
                elif str(entry).strip():
                    certifications.append(
                        {
                            "id": str(uuid.uuid4()),
                            "name": str(entry).strip(),
                            "issuer": "",
                            "date": "",
                        }
                    )
        elif any(word in name for word in ["summary", "profile", "简介", "个人定位"]):
            summary = " ".join(str(item) for item in entries)
    return base_profile(
        file_name,
        first_value(cv, "name"),
        first_value(cv, "email"),
        first_value(cv, "phone"),
        first_value(cv, "location"),
        first_value(cv, "website"),
        first_value(cv, "headline"),
        summary,
        professional_skills,
        experience_entries,
        education_entries,
        projects,
        certifications,
    )


KNOWN_SKILLS = [
    "Python",
    "Java",
    "Golang",
    "Rust",
    "TypeScript",
    "JavaScript",
    "Svelte",
    "React",
    "Vue",
    "FastAPI",
    "Django",
    "Flask",
    "LangChain",
    "RAG",
    "Docker",
    "Kubernetes",
    "Redis",
    "MySQL",
    "PostgreSQL",
    "PyTorch",
    "TensorFlow",
    "AWS",
    "Azure",
]


def profile_from_text(text: str, file_name: str) -> dict[str, Any]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    email_match = re.search(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", text)
    phone_match = re.search(r"(?<!\d)(?:\+?86[- ]?)?1[3-9]\d(?:[- ]?\d){8}(?!\d)", text)
    name = next(
        (line for line in lines[:8] if 1 < len(line) <= 12 and not re.search(r"[@\d:/]", line)), ""
    )
    skills = [
        skill
        for skill in KNOWN_SKILLS
        if re.search(rf"(?<![A-Za-z]){re.escape(skill)}(?![A-Za-z])", text, re.IGNORECASE)
    ]
    headline = next(
        (
            line
            for line in lines[:12]
            if any(
                word in line.lower() for word in ["工程师", "开发", "产品", "designer", "engineer"]
            )
        ),
        "",
    )
    summary = next((line for line in lines if 35 <= len(line) <= 180), "")
    skill_groups = (
        [{"id": str(uuid.uuid4()), "label": "核心技能", "items": skills}] if skills else []
    )
    return base_profile(
        file_name,
        name,
        email_match.group(0) if email_match else "",
        phone_match.group(0) if phone_match else "",
        "",
        "",
        headline,
        summary,
        skill_groups,
        [],
        [],
    )


def base_profile(
    file_name: str,
    name: str,
    email: str,
    phone: str,
    location: str,
    website: str,
    headline: str,
    summary: str,
    professional_skills: list[dict[str, Any]],
    experiences: list[dict[str, Any]],
    education: list[dict[str, Any]],
    projects: list[dict[str, Any]] | None = None,
    certifications: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    if professional_skills and isinstance(professional_skills[0], str):
        professional_skills = [
            {
                "id": str(uuid.uuid4()),
                "label": "核心技能",
                "items": split_skill_items(professional_skills),
            }
        ]
    professional_skills = [group for group in professional_skills if group.get("items")]
    facts = [
        {
            "id": str(uuid.uuid4()),
            "category": "skill",
            "value": skill,
            "source": f"{file_name} · 专业技能 · {group.get('label', '')}",
            "confidence": 0.95,
            "confirmed": False,
        }
        for group in professional_skills
        for skill in group.get("items", [])
    ]
    return {
        "id": "resume-master",
        "name": name,
        "headline": headline,
        "email": email,
        "phone": phone,
        "location": location,
        "website": website,
        "summary": summary,
        "templateId": "ai-engineering",
        "professionalSkills": professional_skills,
        "experiences": experiences,
        "education": education,
        "projects": projects or [],
        "certifications": certifications or [],
        "facts": facts,
        "preferences": {
            "targetRoles": [],
            "cities": [],
            "remotePreference": "flexible",
            "energizingTasks": [],
            "drainingTasks": [],
            "hardConstraints": [],
        },
        "sourceFileName": file_name,
        "updatedAt": now(),
        "version": 1,
    }


def extract_resume(params: dict[str, Any]) -> dict[str, Any]:
    path = pathlib.Path(str(params["path"]))
    file_name = str(params.get("fileName") or path.name)
    text, yaml_data, pages = extract_text(path)
    profile = (
        profile_from_yaml(yaml_data, file_name)
        if isinstance(yaml_data, dict)
        else profile_from_text(text, file_name)
    )
    return {"profile": profile, "rawText": text, "pages": pages}
