from __future__ import annotations

import json
import os
import pathlib
import subprocess
import sys
import tempfile
import types
import unittest
from contextlib import redirect_stdout
from io import StringIO
from unittest.mock import Mock, patch

ROOT = pathlib.Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

import worker  # noqa: E402
from vendor import boss_cdp_raw as boss_vendor  # noqa: E402


class WorkerTests(unittest.TestCase):
    def test_environment_status_reports_missing_chrome_without_downloading(self):
        fake = types.SimpleNamespace(DEFAULT_CHROME_PATH="Z:/missing/google-chrome")
        with patch.object(worker, "load_boss_module", return_value=fake):
            status = worker.environment_status({})
        self.assertEqual(status["protocolVersion"], worker.PROTOCOL_VERSION)
        self.assertFalse(status["chrome"]["installed"])
        self.assertIsNone(status["chrome"]["executablePath"])

    def test_environment_status_reads_version_without_starting_chrome(self):
        with tempfile.TemporaryDirectory() as temporary:
            application = pathlib.Path(temporary) / "Application"
            application.mkdir()
            executable = application / "chrome.exe"
            executable.touch()
            (application / "137.0.1.9").mkdir()
            (application / "138.0.3351.95").mkdir()
            fake = types.SimpleNamespace(DEFAULT_CHROME_PATH=str(executable))
            with (
                patch.object(worker.os, "name", "nt"),
                patch.object(worker.subprocess, "run") as run,
                patch.object(worker, "load_boss_module", return_value=fake),
            ):
                status = worker.environment_status({})
        self.assertTrue(status["chrome"]["installed"])
        self.assertEqual(status["chrome"]["version"], "138.0.3351.95")
        self.assertEqual(status["chrome"]["executablePath"], str(executable))
        run.assert_not_called()

    def test_environment_status_uses_version_command_outside_windows(self):
        with tempfile.TemporaryDirectory() as temporary:
            executable = pathlib.Path(temporary) / "Google Chrome"
            executable.touch()
            fake = types.SimpleNamespace(DEFAULT_CHROME_PATH=str(executable))
            completed = subprocess.CompletedProcess(
                [str(executable), "--version"],
                0,
                stdout="Google Chrome 138.0.7204.168\n",
                stderr="",
            )
            with (
                patch.object(worker.os, "name", "posix"),
                patch.object(worker.subprocess, "run", return_value=completed) as run,
                patch.object(worker, "load_boss_module", return_value=fake),
            ):
                status = worker.environment_status({})
        self.assertTrue(status["chrome"]["installed"])
        self.assertEqual(status["chrome"]["version"], "Google Chrome 138.0.7204.168")
        run.assert_called_once_with(
            [str(executable), "--version"],
            capture_output=True,
            text=True,
            timeout=5,
            check=False,
            encoding="utf-8",
            errors="replace",
        )

    def test_environment_status_tolerates_version_command_failure_outside_windows(self):
        with tempfile.TemporaryDirectory() as temporary:
            executable = pathlib.Path(temporary) / "Google Chrome"
            executable.touch()
            fake = types.SimpleNamespace(DEFAULT_CHROME_PATH=str(executable))
            with (
                patch.object(worker.os, "name", "posix"),
                patch.object(
                    worker.subprocess,
                    "run",
                    side_effect=subprocess.TimeoutExpired([str(executable), "--version"], 5),
                ),
                patch.object(worker, "load_boss_module", return_value=fake),
            ):
                status = worker.environment_status({})
        self.assertTrue(status["chrome"]["installed"])
        self.assertIsNone(status["chrome"]["version"])

    def test_clear_boss_data_requires_no_remaining_profile_processes(self):
        with tempfile.TemporaryDirectory() as temporary:
            profile = pathlib.Path(temporary) / "profile"
            result = pathlib.Path(temporary) / "result"
            profile.mkdir()
            result.mkdir()
            fake = types.SimpleNamespace(
                DEFAULT_CDP_DATA_DIR=str(profile),
                DEFAULT_RESULT_DIR=str(result),
                chrome_pids_for_user_data_dir=Mock(return_value=[]),
            )
            with (
                patch.object(worker, "load_boss_module", return_value=fake),
                patch.object(worker, "close_boss_session", return_value={"cleanupSucceeded": True}),
            ):
                outcome = worker.clear_boss_data({})
            self.assertEqual(outcome["remainingPids"], [])
            self.assertFalse(profile.exists())
            self.assertFalse(result.exists())

    def test_cdp_errors_are_not_converted_to_empty_values(self):
        class FakeSocket:
            def send(self, _message):
                pass

            def recv(self):
                return json.dumps({"id": 1, "error": {"code": -32000, "message": "detached"}})

        session = boss_vendor.CDPSession.__new__(boss_vendor.CDPSession)
        session.mid = 0
        session.ws = FakeSocket()
        with self.assertRaisesRegex(RuntimeError, "detached"):
            session.send("Runtime.evaluate")

        session.send = lambda *_args, **_kwargs: {
            "result": {"exceptionDetails": {"text": "ReferenceError"}}
        }
        with self.assertRaisesRegex(RuntimeError, "ReferenceError"):
            session.eval_js("missing", "session")

    def test_cdp_websocket_suppresses_origin_header(self):
        response = types.SimpleNamespace(
            json=lambda: {"webSocketDebuggerUrl": "ws://127.0.0.1/devtools/browser/test"}
        )
        requests_module = types.SimpleNamespace(get=Mock(return_value=response))
        connect = Mock(return_value=types.SimpleNamespace())
        websocket_module = types.SimpleNamespace(create_connection=connect)
        with (
            patch.object(boss_vendor, "require_runtime_dependencies", return_value=True),
            patch.object(boss_vendor, "requests", requests_module),
            patch.object(boss_vendor, "websocket", websocket_module),
        ):
            boss_vendor.CDPSession(9222)
        self.assertTrue(connect.call_args.kwargs["suppress_origin"])

    def test_unknown_city_resolution_never_uses_live_network_map(self):
        self.assertFalse(hasattr(boss_vendor, "load_live_city_maps"))
        self.assertEqual(boss_vendor.resolve_city("不存在城市"), ("不存在城市", "不存在城市"))

    def test_pdf_render_closes_every_resource_when_save_fails(self):
        closed = []

        class Resource:
            def __init__(self, name):
                self.name = name

            def close(self):
                closed.append(self.name)

        class Image(Resource):
            def save(self, *_args, **_kwargs):
                raise RuntimeError("save failed")

        class Bitmap(Resource):
            def to_pil(self):
                return Image("image")

        class Page(Resource):
            def render(self, **_kwargs):
                return Bitmap("bitmap")

        class Document(Resource):
            def __getitem__(self, _index):
                return Page("page")

        fake_pdfium = types.SimpleNamespace(PdfDocument=lambda _path: Document("document"))
        with (
            patch.dict(sys.modules, {"pypdfium2": fake_pdfium}),
            tempfile.TemporaryDirectory() as temporary,
        ):
            with self.assertRaisesRegex(RuntimeError, "save failed"):
                worker.render_pdf_page(pathlib.Path("resume.pdf"), 0, pathlib.Path(temporary))
        self.assertEqual(closed, ["image", "bitmap", "page", "document"])

    def test_sigterm_cleanup_is_scoped_to_active_boss(self):
        fake = object()
        worker._active_boss = fake
        worker._cleaning_boss = False
        with patch.object(worker, "close_boss_session") as close:
            worker.cleanup_active_boss()
        close.assert_called_once()
        self.assertIsNone(worker._active_boss)

    def test_text_extraction_does_not_invent_experience(self):
        profile = worker.profile_from_text(
            "林知远\nAI 应用研发工程师\nlin@example.com\n熟悉 Python、FastAPI、RAG 和 Docker。",
            "resume.txt",
        )
        self.assertEqual(profile["name"], "林知远")
        self.assertIn("Python", profile["professionalSkills"][0]["items"])
        self.assertEqual(profile["experiences"], [])

    def test_rendercv_yaml_is_structured(self):
        profile = worker.base_profile(
            "resume.yaml",
            "林知远",
            "lin@example.com",
            "",
            "上海",
            "",
            "AI 工程师",
            "简介",
            ["Python"],
            [],
            [],
        )
        data = worker.profile_to_rendercv(profile)
        self.assertEqual(data["cv"]["name"], "林知远")
        self.assertEqual(
            data["cv"]["sections"]["专业技能"], [{"label": "核心技能", "details": "Python"}]
        )

    def test_role_templates_keep_rendercv_section_order(self):
        profile = worker.base_profile(
            "resume.yaml",
            "示例候选人",
            "sample@example.invalid",
            "",
            "上海",
            "",
            "候选人",
            "简介",
            ["SQL"],
            [],
            [],
        )
        profile["experiences"] = [
            {
                "company": "示例公司",
                "position": "示例岗位",
                "location": "上海",
                "startDate": "2022.01",
                "endDate": "至今",
                "highlights": ["示例成果"],
            }
        ]
        profile["projects"] = [
            {
                "name": "示例项目",
                "summary": "示例简介",
                "startDate": "2024.01",
                "endDate": "2024.06",
                "highlights": ["示例成果"],
            }
        ]
        profile["certifications"] = [{"name": "示例证书", "issuer": "示例机构", "date": "2023.01"}]
        profile["education"] = [
            {
                "institution": "示例大学",
                "area": "示例专业",
                "degree": "本科",
                "startDate": "2018.09",
                "endDate": "2022.06",
                "highlights": [],
            }
        ]

        profile["templateId"] = "data-analysis"
        data_sections = list(worker.profile_to_rendercv(profile)["cv"]["sections"])
        self.assertEqual(data_sections[:4], ["个人简介", "专业技能", "工作经历", "项目经历"])

        profile["templateId"] = "finance-accounting"
        finance_sections = list(worker.profile_to_rendercv(profile)["cv"]["sections"])
        self.assertEqual(
            finance_sections[:4], ["个人简介", "工作经历", "证书 / 专业资质", "专业技能"]
        )

    def test_scrape_rejects_blank_keyword_without_loading_boss(self):
        with patch.object(worker, "load_boss_module") as load_boss:
            with self.assertRaisesRegex(ValueError, "岗位关键词不能为空"):
                worker.scrape_jobs({"keyword": "   ", "city": "上海"})
        load_boss.assert_not_called()

    def test_rendercv_yaml_keeps_education_separate(self):
        import yaml

        path = ROOT / "tests" / "fixtures" / "sample_resume.yaml"
        profile = worker.profile_from_yaml(
            yaml.safe_load(path.read_text(encoding="utf-8")), path.name
        )
        self.assertEqual(len(profile["experiences"]), 1)
        self.assertEqual(len(profile["education"]), 1)
        self.assertEqual(profile["education"][0]["institution"], "浙江工业大学")

    def test_date_ranges_are_normalized_and_single_dates_have_no_prefix(self):
        self.assertEqual(worker.normalize_date_pair("", "2024.12 - 至今"), ("2024.12", "至今"))
        self.assertEqual(worker.normalize_date_pair("2019.09–2023.06", ""), ("2019.09", "2023.06"))
        self.assertEqual(
            worker.render_date_fields({"startDate": "", "endDate": "2024.12"}), ("2024-12", None)
        )
        self.assertEqual(worker.rendercv_date("2024.12"), "2024-12")
        self.assertEqual(worker.rendercv_date("至今", end_date=True), "present")

    def test_color_themes_share_one_layout_and_only_change_accent_colors(self):
        profile = worker.base_profile("resume.pdf", "测试", "", "", "", "", "", "", [], [], [])
        expected = {
            "pine": "#176B57",
            "navy": "#1F407A",
            "graphite": "#24292F",
        }
        layouts = []
        for color_theme, accent in expected.items():
            design = worker.profile_to_rendercv(profile, color_theme)["design"]
            self.assertEqual(design["theme"], "classic")
            self.assertEqual(design["colors"]["name"], accent)
            self.assertEqual(design["colors"]["section_titles"], accent)
            neutral_colors = {**design["colors"]}
            for key in ("name", "headline", "connections", "section_titles", "links"):
                neutral_colors[key] = None
            layouts.append({**design, "colors": neutral_colors})
        self.assertTrue(all(layout == layouts[0] for layout in layouts[1:]))

    def test_unknown_resume_color_theme_is_rejected(self):
        profile = worker.base_profile("resume.pdf", "测试", "", "", "", "", "", "", [], [], [])
        with self.assertRaisesRegex(ValueError, "不支持的简历颜色主题"):
            worker.profile_to_rendercv(profile, "unknown")

    def test_resume_design_matches_the_reference_rendercv_style(self):
        profile = worker.base_profile("resume.pdf", "测试", "", "", "", "", "", "", [], [], [])
        data = worker.profile_to_rendercv(profile)
        design = data["design"]

        self.assertEqual(design["page"]["top_margin"], "1.2cm")
        self.assertEqual(design["page"]["left_margin"], "1.35cm")
        self.assertEqual(design["typography"]["font_family"]["body"], "Microsoft YaHei")
        self.assertEqual(design["header"]["alignment"], "center")
        self.assertFalse(design["header"]["connections"]["show_icons"])
        self.assertEqual(design["header"]["connections"]["separator"], "|")
        self.assertEqual(design["entries"]["date_and_location_width"], "4.6cm")
        self.assertEqual(
            design["templates"]["experience_entry"]["date_and_location_column"], "LOCATION · DATE"
        )
        self.assertIn("Dify", data["settings"]["bold_keywords"])

    def test_all_color_themes_generate_pdf_files(self):
        from rendercv.renderer import pdf_png

        bundled_library = pathlib.Path(pdf_png.__file__).parent / "rendercv_typst" / "lib.typ"
        bundled_before = bundled_library.read_bytes()
        profile = worker.base_profile(
            "resume.pdf",
            "Candidate",
            "candidate@example.com",
            "",
            "Shanghai",
            "",
            "Engineer",
            "Summary",
            [{"id": "skills", "label": "Core", "items": ["Python"]}],
            [
                {
                    "id": "experience",
                    "company": "Example",
                    "position": "Engineer",
                    "location": "Shanghai",
                    "startDate": "",
                    "endDate": "2024.12 - present",
                    "highlights": ["Delivered project"],
                }
            ],
            [
                {
                    "id": "education",
                    "institution": "Example University",
                    "area": "Computer Science",
                    "degree": "Bachelor",
                    "degreeDetail": "",
                    "startDate": "2018.09",
                    "endDate": "2022.06",
                    "highlights": [],
                }
            ],
        )
        with tempfile.TemporaryDirectory() as temporary:
            for color_theme in worker.RESUME_COLOR_THEMES:
                result = worker.render_resume(
                    {
                        "profile": profile,
                        "colorTheme": color_theme,
                        "outputPath": str(pathlib.Path(temporary) / f"{color_theme}.pdf"),
                    }
                )
                output = pathlib.Path(result["path"])
                self.assertTrue(output.exists())
                self.assertGreater(output.stat().st_size, 1_000)
            self.assertTrue(
                all(path.suffix == ".pdf" for path in pathlib.Path(temporary).iterdir())
            )
        self.assertEqual(bundled_library.read_bytes(), bundled_before)

    def test_docx_extraction_includes_body_tables_headers_and_footers(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as temporary:
            path = pathlib.Path(temporary) / "resume.docx"
            document = Document()
            document.sections[0].header.paragraphs[0].text = "页眉联系方式"
            document.add_paragraph("张三")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "教育经历"
            table.cell(0, 1).text = "示例大学"
            document.sections[0].footer.paragraphs[0].text = "个人网站"
            document.save(path)

            text = worker.extract_docx_text(path)
            self.assertIn("张三", text)
            self.assertIn("教育经历 | 示例大学", text)
            self.assertIn("页眉联系方式", text)
            self.assertIn("个人网站", text)

    def test_scanned_pdf_pages_are_rendered_for_vision_transcription(self):
        from PIL import Image, ImageDraw

        with tempfile.TemporaryDirectory() as temporary:
            path = pathlib.Path(temporary) / "scan.pdf"
            image = Image.new("RGB", (900, 1200), "white")
            ImageDraw.Draw(image).text((80, 80), "SCANNED RESUME", fill="black")
            image.save(path, "PDF")
            image.close()

            text, yaml_data, pages = worker.extract_text(path)
            self.assertEqual(text, "")
            self.assertIsNone(yaml_data)
            self.assertEqual(len(pages), 1)
            self.assertTrue(pathlib.Path(pages[0]["imagePath"]).exists())

    def test_yaml_preserves_grouped_skills_positioning_projects_and_dates(self):
        profile = worker.profile_from_yaml(
            {
                "cv": {
                    "name": "林知远",
                    "sections": {
                        "个人定位": ["面向 AI 工程岗位。"],
                        "专业技能": [{"label": "后端与数据", "details": "Python, PostgreSQL"}],
                        "工作经历": [
                            {"company": "示例公司", "position": "工程师", "date": "2024.01 - 至今"}
                        ],
                        "项目经历": [
                            {
                                "name": "RAG 平台",
                                "summary": "本地知识库",
                                "highlights": ["完成上线"],
                            }
                        ],
                    },
                }
            },
            "resume.yaml",
        )
        self.assertEqual(profile["summary"], "面向 AI 工程岗位。")
        self.assertEqual(profile["professionalSkills"][0]["label"], "后端与数据")
        self.assertEqual(profile["professionalSkills"][0]["items"], ["Python", "PostgreSQL"])
        self.assertEqual(profile["experiences"][0]["startDate"], "2024.01")
        self.assertEqual(profile["experiences"][0]["endDate"], "至今")
        self.assertEqual(profile["projects"][0]["name"], "RAG 平台")

    def test_market_report_stays_scoped(self):
        report = worker.market_report(
            [{"skills": ["Python"], "experience": "3-5年", "degree": "本科"}], "AI Agent", "上海"
        )
        self.assertIn("本次岗位样本观察", report)
        self.assertIn("有限页样本", report)
        self.assertIn("不能作为候选人经历证据", report)
        self.assertIn("上海", report)

    def test_shanghai_timezone_is_explicit(self):
        self.assertTrue(worker.now().endswith("+08:00"))

    def test_setup_boss_reset_is_safe_and_returns_structured_outcome(self):
        class FakeBoss:
            DEFAULT_CDP_DATA_DIR = "fake-profile"
            reset_values = []
            stop_calls = 0

            @classmethod
            def run_setup_chrome(cls, *_args, **kwargs):
                cls.reset_values.append(kwargs["reset_profile"])
                return 0

            @classmethod
            def stop_cdp_chrome(cls, data_dir):
                self.assertEqual(data_dir, cls.DEFAULT_CDP_DATA_DIR)
                cls.stop_calls += 1
                return 1

            @staticmethod
            def chrome_pids_for_user_data_dir(_data_dir):
                return []

        with (
            patch.object(worker, "load_boss_module", return_value=FakeBoss()),
            patch.object(worker, "emit"),
        ):
            result = worker.setup_boss({"resetProfile": True, "loginTimeout": 10})

        self.assertEqual(FakeBoss.reset_values, [True])
        self.assertEqual(FakeBoss.stop_calls, 2)
        self.assertEqual(
            result,
            {
                "loginSucceeded": True,
                "resetRequested": True,
                "cleanupSucceeded": True,
                "closedProcesses": 2,
                "error": None,
            },
        )

    def test_setup_boss_failure_still_closes_dedicated_chrome(self):
        class FakeBoss:
            DEFAULT_CDP_DATA_DIR = "fake-profile"
            stop_calls = 0

            @staticmethod
            def run_setup_chrome(*_args, **_kwargs):
                print("登录等待超时")
                return 1

            @classmethod
            def stop_cdp_chrome(cls, _data_dir):
                cls.stop_calls += 1
                return 1

            @staticmethod
            def chrome_pids_for_user_data_dir(_data_dir):
                return []

        with (
            patch.object(worker, "load_boss_module", return_value=FakeBoss()),
            patch.object(worker, "emit"),
        ):
            result = worker.setup_boss({"loginTimeout": 1})

        self.assertFalse(result["loginSucceeded"])
        self.assertFalse(result["resetRequested"])
        self.assertTrue(result["cleanupSucceeded"])
        self.assertEqual(result["closedProcesses"], 1)
        self.assertIn("登录等待超时", result["error"])
        self.assertEqual(FakeBoss.stop_calls, 1)

    def test_setup_boss_does_not_reset_if_profile_process_survives(self):
        class FakeBoss:
            DEFAULT_CDP_DATA_DIR = "fake-profile"
            setup_called = False

            @classmethod
            def run_setup_chrome(cls, *_args, **_kwargs):
                cls.setup_called = True
                return 0

            @staticmethod
            def stop_cdp_chrome(_data_dir):
                return 0

            @staticmethod
            def chrome_pids_for_user_data_dir(_data_dir):
                return [91]

        with (
            patch.object(worker, "load_boss_module", return_value=FakeBoss()),
            patch.object(worker, "emit"),
        ):
            result = worker.setup_boss({"resetProfile": True})

        self.assertFalse(FakeBoss.setup_called)
        self.assertFalse(result["loginSucceeded"])
        self.assertFalse(result["cleanupSucceeded"])
        self.assertIn("91", result["error"])

    def test_close_boss_operation_returns_verified_cleanup(self):
        class FakeBoss:
            DEFAULT_CDP_DATA_DIR = "fake-profile"

            @staticmethod
            def stop_cdp_chrome(_data_dir):
                return 2

            @staticmethod
            def chrome_pids_for_user_data_dir(_data_dir):
                return []

        with (
            patch.object(worker, "load_boss_module", return_value=FakeBoss()),
            patch.object(worker, "emit"),
        ):
            result = worker.OPERATIONS["close_boss"]({})

        self.assertEqual(
            result,
            {
                "cleanupSucceeded": True,
                "closedProcesses": 2,
                "error": None,
            },
        )

    def test_profile_scoped_stop_never_terminates_normal_chrome(self):
        dedicated = r"C:\tmp\boss-profile"
        regular = r"C:\Users\kw\AppData\Local\Google\Chrome\User Data"
        running = [
            (11, f'chrome.exe --user-data-dir="{regular}"'),
            (22, f'chrome.exe --user-data-dir="{dedicated}" --remote-debugging-port=9222'),
        ]

        def terminate(pid, force=False):
            self.assertFalse(force)
            running[:] = [item for item in running if item[0] != pid]

        with (
            patch.object(
                boss_vendor, "iter_chrome_process_commands", side_effect=lambda: list(running)
            ),
            patch.object(boss_vendor, "terminate_process", side_effect=terminate) as terminate_mock,
            patch.object(boss_vendor.time, "sleep"),
        ):
            stopped = boss_vendor.stop_cdp_chrome(dedicated)

        self.assertEqual(stopped, 1)
        self.assertEqual(terminate_mock.call_count, 1)
        self.assertEqual(terminate_mock.call_args.args, (22,))
        self.assertEqual(running[0][0], 11)

    def test_windows_chrome_process_lookup_hides_powershell_window(self):
        completed = subprocess.CompletedProcess(args=[], returncode=0, stdout="[]", stderr="")
        with (
            patch.object(boss_vendor.platform, "system", return_value="Windows"),
            patch.object(boss_vendor.subprocess, "run", return_value=completed) as run,
        ):
            self.assertEqual(boss_vendor.iter_chrome_process_commands(), [])

        self.assertEqual(
            run.call_args.kwargs["creationflags"],
            getattr(subprocess, "CREATE_NO_WINDOW", 0x08000000),
        )

    def test_windows_chrome_termination_hides_taskkill_window(self):
        with (
            patch.object(boss_vendor.platform, "system", return_value="Windows"),
            patch.object(boss_vendor.subprocess, "run") as run,
        ):
            boss_vendor.terminate_process(22, force=True)

        self.assertEqual(run.call_args.args[0], ["taskkill", "/PID", "22", "/T", "/F"])
        self.assertEqual(
            run.call_args.kwargs["creationflags"],
            getattr(subprocess, "CREATE_NO_WINDOW", 0x08000000),
        )

    def test_profile_reset_refuses_to_delete_while_pid_remains(self):
        with (
            patch.object(boss_vendor, "stop_cdp_chrome", return_value=0),
            patch.object(boss_vendor, "chrome_pids_for_user_data_dir", return_value=[22]),
            patch.object(boss_vendor.os.path, "exists", return_value=True),
            patch.object(boss_vendor.shutil, "rmtree") as rmtree,
        ):
            with self.assertRaisesRegex(RuntimeError, "remaining PIDs"):
                boss_vendor.prepare_cdp_profile(reset=True)

        rmtree.assert_not_called()

    def test_scrape_resolves_city_code_and_connects_boss(self):
        call_order = []

        class FakeBoss:
            DEFAULT_CDP_DATA_DIR = "fake-profile"
            setup_called = False
            stop_called = False
            received_city = ""
            received_pages = 0

            @staticmethod
            def resolve_city(city):
                self.assertEqual(city, "上海")
                return "上海", "101020100"

            @classmethod
            def run_setup_chrome(cls, *_args, **kwargs):
                cls.setup_called = True
                self.assertFalse(kwargs["reset_profile"])
                self.assertTrue(kwargs["wait_login"])
                self.assertEqual(kwargs["login_timeout"], 300)
                call_order.append("setup")
                return 0

            @classmethod
            def scrape_list(cls, _keyword, city, pages, _filters, _output, **_kwargs):
                self.assertEqual(call_order, ["setup"])
                call_order.append("list")
                cls.received_city = city
                cls.received_pages = pages
                result = {
                    "jobs": [
                        {
                            "job_id": "job-1",
                            "title": "AI 工程师",
                            "boss_name": "示例公司",
                            "location": "上海·浦东新区",
                            "salary": "20-30K",
                        }
                    ]
                }
                _kwargs["on_job"](result["jobs"][0])
                return result

            @staticmethod
            def scrape_details(listing, *_args, **_kwargs):
                detail = {"job_id": "job-1", "jd": "负责 AI 应用研发", "skill_tags": ["Python"]}
                _kwargs["on_detail"](listing["jobs"][0], detail)
                _kwargs["on_progress"](
                    status="success",
                    processed=1,
                    total=1,
                    succeeded=1,
                    skipped=0,
                    failed=0,
                    title="AI 工程师",
                    message="详情抓取成功",
                )
                return [detail]

            @classmethod
            def stop_cdp_chrome(cls, data_dir):
                self.assertEqual(data_dir, cls.DEFAULT_CDP_DATA_DIR)
                cls.stop_called = True
                return 1

            @staticmethod
            def chrome_pids_for_user_data_dir(_data_dir):
                return []

        fake = FakeBoss()
        events = []
        with (
            patch.object(worker, "load_boss_module", return_value=fake),
            patch.object(
                worker, "emit", side_effect=lambda kind, **payload: events.append((kind, payload))
            ),
            redirect_stdout(StringIO()),
        ):
            result = worker.scrape_jobs({"keyword": "AI Agent", "city": " 上海 "})

        self.assertTrue(fake.setup_called)
        self.assertEqual(call_order, ["setup", "list"])
        self.assertTrue(fake.stop_called)
        self.assertEqual(fake.received_city, "101020100")
        self.assertEqual(fake.received_pages, 1)
        self.assertEqual(result["resolvedCity"], "上海")
        self.assertEqual(result["cityCode"], "101020100")
        self.assertEqual(len(result["jobs"]), 1)
        job_events = [payload["job"] for kind, payload in events if kind == "job"]
        self.assertEqual(len(job_events), 2)
        self.assertEqual(job_events[0]["description"], "")
        self.assertEqual(job_events[1]["description"], "负责 AI 应用研发")
        self.assertEqual(
            [payload["phase"] for kind, payload in events if kind == "job"],
            ["list", "detail"],
        )

    def test_scrape_caps_pages_and_filters_only_detail_candidates(self):
        class FakeBoss:
            DEFAULT_CDP_DATA_DIR = "fake-profile"
            received_pages = 0
            detail_job_ids = []

            @staticmethod
            def resolve_city(_city):
                return "上海", "101020100"

            @staticmethod
            def run_setup_chrome(*_args, **_kwargs):
                return 0

            @classmethod
            def scrape_list(cls, _keyword, _city, pages, _filters, _output, **kwargs):
                cls.received_pages = pages
                jobs = [
                    {"job_id": "job-done", "title": "已抓岗位", "boss_name": "甲公司"},
                    {"job_id": "job-new", "title": "新岗位", "boss_name": "乙公司"},
                ]
                for job in jobs:
                    kwargs["on_job"](job)
                return {"jobs": jobs, "total": len(jobs)}

            @classmethod
            def scrape_details(cls, listing, *_args, **kwargs):
                cls.detail_job_ids = [job["job_id"] for job in listing["jobs"]]
                raw = listing["jobs"][0]
                detail = {"job_id": raw["job_id"], "jd": "有效 JD", "skill_tags": []}
                kwargs["on_detail"](raw, detail)
                kwargs["on_progress"](
                    status="success",
                    processed=1,
                    total=1,
                    succeeded=1,
                    skipped=0,
                    failed=0,
                    title=raw["title"],
                    message="详情抓取成功",
                )
                return [detail]

            @staticmethod
            def stop_cdp_chrome(_data_dir):
                return 0

            @staticmethod
            def chrome_pids_for_user_data_dir(_data_dir):
                return []

        events = []
        with (
            patch.object(worker, "load_boss_module", return_value=FakeBoss()),
            patch.object(
                worker, "emit", side_effect=lambda kind, **payload: events.append((kind, payload))
            ),
        ):
            result = worker.scrape_jobs(
                {
                    "keyword": "AI",
                    "city": "上海",
                    "pages": 99,
                    "completedDetailExternalIds": ["job-done"],
                }
            )

        self.assertEqual(FakeBoss.received_pages, 5)
        self.assertEqual(FakeBoss.detail_job_ids, ["job-new"])
        self.assertEqual(
            [payload["phase"] for kind, payload in events if kind == "job"],
            [
                "list",
                "list",
                "detail",
            ],
        )
        self.assertEqual(
            result["detailSummary"],
            {
                "total": 2,
                "succeeded": 1,
                "skipped": 1,
                "failed": 0,
                "processed": 2,
            },
        )
        final_progress = [
            payload for kind, payload in events if kind == "progress" and payload["progress"] == 78
        ]
        self.assertEqual(final_progress[-1]["detailSkipped"], 1)
        self.assertEqual(final_progress[-1]["detailFailed"], 0)

    def test_vendor_limits_pages_and_has_required_offline_cities(self):
        self.assertEqual(boss_vendor.MAX_PAGES, 5)
        self.assertGreaterEqual(len(boss_vendor.CITY_MAP), 300)
        self.assertEqual(boss_vendor.CITY_MAP["昆明"], "101290100")
        self.assertEqual(boss_vendor.CITY_MAP["南昌"], "101240100")
        self.assertEqual(boss_vendor.CITY_MAP["石家庄"], "101090100")
        self.assertEqual(boss_vendor.resolve_city("赣州"), ("赣州", "101240700"))
        self.assertEqual(boss_vendor.resolve_city("101240700"), ("赣州", "101240700"))

    def test_detail_extractor_removes_page_chrome_and_recruiter_footer(self):
        description = "负责 AI 产品规划、需求分析、研发协作和上线复盘。\n" * 8
        page_text = (
            "微信扫码分享 举报\n职位描述\n"
            f"{description}"
            "张女士\n今日活跃\n示例公司\n·\n招聘者\n竞争力分析\n"
            "查看完整个人竞争力\nBOSS 安全提示\n公司工商信息\n更多职位"
        )

        jd = boss_vendor.extract_job_description({"jd": page_text, "page_text": page_text})

        self.assertEqual(jd, description.strip())
        self.assertNotIn("张女士", jd)
        self.assertNotIn("BOSS 安全提示", jd)
        self.assertNotIn("jd = body.substring", boss_vendor.EXTRACT_DETAIL_JS)
        self.assertNotIn("pageText.substring", boss_vendor.EXTRACT_DETAIL_JS)
        self.assertIn("page_text", boss_vendor.EXTRACT_DETAIL_JS)

    def test_detail_extractor_appends_company_introduction(self):
        description = "负责 AI 产品规划、需求分析、研发协作和上线复盘。\n" * 8
        page_text = (
            "职位描述\n"
            f"{description}"
            "张女士\n今日活跃\n示例公司\n·\n招聘者\n竞争力分析\n"
            "公司介绍\n示例公司专注于企业级人工智能产品。\n"
            "拥有成熟的研发团队和多个落地案例。\n点击查看地图\n公司工商信息"
        )

        jd = boss_vendor.extract_job_description({"jd": page_text, "page_text": page_text})

        self.assertEqual(
            jd,
            description.strip() + "\n\n公司介绍\n示例公司专注于企业级人工智能产品。\n"
            "拥有成熟的研发团队和多个落地案例。",
        )
        self.assertNotIn("点击查看地图", jd)
        self.assertNotIn("公司工商信息", jd)

    def test_detail_extractor_uses_more_jobs_as_company_introduction_boundary(self):
        description = "负责 AI 平台研发、模型部署和业务场景落地。\n" * 8
        page_text = (
            f"职位描述\n{description}BOSS 安全提示\n"
            "公司介绍\n这是一段公司介绍。\n更多职位\n推荐岗位"
        )

        jd = boss_vendor.extract_job_description({"jd": page_text, "page_text": page_text})

        self.assertTrue(jd.endswith("\n\n公司介绍\n这是一段公司介绍。"))
        self.assertNotIn("更多职位", jd)

    def test_detail_extractor_does_not_append_without_company_introduction(self):
        description = "负责 AI 产品规划、需求分析、研发协作和上线复盘。\n" * 8
        page_text = f"职位描述\n{description}BOSS 安全提示\n更多职位"

        jd = boss_vendor.extract_job_description({"jd": page_text, "page_text": page_text})

        self.assertEqual(jd, description.strip())

    def test_detail_extractor_rejects_navigation_and_short_pages(self):
        with self.assertRaisesRegex(boss_vendor.DetailExtractionError, "navigation chrome"):
            boss_vendor.extract_job_description(
                {
                    "jd": "",
                    "page_text": "首页\n职位\n公司\n校园\n无障碍专区\n热门职位",
                }
            )
        with self.assertRaisesRegex(boss_vendor.DetailExtractionError, "too short"):
            boss_vendor.extract_job_description({"jd": "职位描述\n只有一句话"})

    def test_detail_failures_and_empty_jd_do_not_stop_later_jobs(self):
        class FakeSession:
            instances = []

            def __init__(self, _port):
                self.index = len(self.instances)
                self.closed = False
                self.closed_target = False
                self.calls = []
                self.instances.append(self)

            def send(self, method, params, session_id=None):
                self.calls.append((method, params, session_id))
                if method == "Target.createTarget":
                    if self.index == 0:
                        raise RuntimeError("首条导航失败")
                    return {"result": {"targetId": f"target-{self.index}"}}
                if method == "Target.attachToTarget":
                    return {"result": {"sessionId": f"session-{self.index}"}}
                if method == "Target.closeTarget":
                    self.closed_target = True
                return {"result": {}}

            def eval_js(self, script, _session_id=None):
                if script != boss_vendor.EXTRACT_DETAIL_JS:
                    return None
                if self.index == 1:
                    return json.dumps({"jd": "   ", "tags": []})
                return json.dumps(
                    {
                        "jd": "职位描述\n" + "负责 AI 平台研发、模型部署和业务场景落地。\n" * 8,
                        "tags": ["Python"],
                    }
                )

            def close(self):
                self.closed = True

        jobs = [
            {
                "job_id": f"job-{index}",
                "title": f"岗位 {index}",
                "boss_name": "示例公司",
                "job_link": f"https://www.zhipin.com/job_detail/job-{index}.html",
            }
            for index in range(3)
        ]
        details = []
        progress = []
        with (
            tempfile.TemporaryDirectory() as temporary,
            patch.object(boss_vendor, "CDPSession", FakeSession),
            patch.object(boss_vendor, "incr_request"),
            patch.object(boss_vendor.time, "sleep"),
            patch.object(boss_vendor.random, "uniform", return_value=0),
            patch.object(boss_vendor.random, "randint", return_value=0),
            patch.object(boss_vendor.random, "random", return_value=1),
        ):
            result = boss_vendor.scrape_details(
                {"jobs": jobs},
                output_path=str(pathlib.Path(temporary) / "details.json"),
                on_detail=lambda raw, detail: details.append((raw, detail)),
                on_progress=lambda **payload: progress.append(payload),
            )

        self.assertEqual([detail["job_id"] for detail in result], ["job-2"])
        self.assertEqual(len(details), 1)
        self.assertEqual(
            [payload["status"] for payload in progress], ["failed", "failed", "success"]
        )
        self.assertEqual(progress[-1]["succeeded"], 1)
        self.assertEqual(progress[-1]["failed"], 2)
        self.assertTrue(all(session.closed for session in FakeSession.instances))
        self.assertFalse(FakeSession.instances[0].closed_target)
        self.assertTrue(FakeSession.instances[1].closed_target)
        self.assertTrue(FakeSession.instances[2].closed_target)
        calls = FakeSession.instances[2].calls
        self.assertIn(
            ("Target.createTarget", {"url": "about:blank", "background": True}, None), calls
        )
        injection_index = next(
            index
            for index, call in enumerate(calls)
            if call[0] == "Page.addScriptToEvaluateOnNewDocument"
        )
        navigation_index = next(
            index for index, call in enumerate(calls) if call[0] == "Page.navigate"
        )
        self.assertLess(injection_index, navigation_index)
        self.assertIn("visibilityState", calls[injection_index][1]["source"])

    def test_detail_login_wall_stops_run_and_closes_target(self):
        class FakeSession:
            instance = None

            def __init__(self, _port):
                self.closed = False
                self.closed_target = False
                FakeSession.instance = self

            def send(self, method, _params, _session_id=None):
                if method == "Target.createTarget":
                    return {"result": {"targetId": "target-login"}}
                if method == "Target.attachToTarget":
                    return {"result": {"sessionId": "session-login"}}
                if method == "Target.closeTarget":
                    self.closed_target = True
                return {"result": {}}

            def eval_js(self, script, _session_id=None):
                if script == boss_vendor.EXTRACT_DETAIL_JS:
                    return json.dumps(
                        {
                            "jd": "",
                            "page_text": "职位描述\n负责产品规划\n登录查看完整内容",
                            "tags": [],
                        }
                    )
                return None

            def close(self):
                self.closed = True

        job = {
            "job_id": "blocked",
            "title": "AI 产品经理",
            "boss_name": "示例公司",
            "job_link": "https://www.zhipin.com/job_detail/blocked.html",
        }
        progress = []
        with (
            tempfile.TemporaryDirectory() as temporary,
            patch.object(boss_vendor, "CDPSession", FakeSession),
            patch.object(boss_vendor, "incr_request"),
            patch.object(boss_vendor.time, "sleep"),
            patch.object(boss_vendor.random, "uniform", return_value=0),
            patch.object(boss_vendor.random, "randint", return_value=0),
            patch.object(boss_vendor.random, "random", return_value=1),
        ):
            output = pathlib.Path(temporary) / "details.json"
            with self.assertRaisesRegex(RuntimeError, "登录状态已失效"):
                boss_vendor.scrape_details(
                    {"jobs": [job]},
                    output_path=str(output),
                    on_progress=lambda **payload: progress.append(payload),
                )
            self.assertFalse(output.exists())

        self.assertEqual(progress[-1]["status"], "failed")
        self.assertTrue(FakeSession.instance.closed_target)
        self.assertTrue(FakeSession.instance.closed)

    def test_scrape_exception_still_closes_dedicated_chrome(self):
        class FakeBoss:
            DEFAULT_CDP_DATA_DIR = "fake-profile"
            stop_called = False

            @staticmethod
            def resolve_city(_city):
                return "上海", "101020100"

            @staticmethod
            def run_setup_chrome(*_args, **_kwargs):
                return 0

            @staticmethod
            def scrape_list(*_args, **_kwargs):
                raise RuntimeError("验证码拦截")

            @classmethod
            def stop_cdp_chrome(cls, data_dir):
                self.assertEqual(data_dir, cls.DEFAULT_CDP_DATA_DIR)
                cls.stop_called = True
                return 1

            @staticmethod
            def chrome_pids_for_user_data_dir(_data_dir):
                return []

        with (
            patch.object(worker, "load_boss_module", return_value=FakeBoss()),
            patch.object(worker, "emit"),
        ):
            with self.assertRaisesRegex(RuntimeError, "验证码拦截"):
                worker.scrape_jobs({"keyword": "AI", "city": "上海", "pages": 1})

        self.assertTrue(FakeBoss.stop_called)

    def test_unknown_city_is_rejected_before_browser_connection(self):
        class FakeBoss:
            DEFAULT_CDP_DATA_DIR = "fake-profile"

            @staticmethod
            def resolve_city(city):
                return city, city

            @staticmethod
            def stop_cdp_chrome(_data_dir):
                return 0

            @staticmethod
            def chrome_pids_for_user_data_dir(_data_dir):
                return []

        with (
            patch.object(worker, "load_boss_module", return_value=FakeBoss()),
            patch.object(worker, "emit"),
        ):
            with self.assertRaisesRegex(RuntimeError, "无法识别城市"):
                worker.scrape_jobs({"keyword": "AI", "city": "乱码城市", "pages": 1})

    def test_jsonl_process_boundary_forces_utf8(self):
        fixture = ROOT / "tests" / "fixtures" / "sample_resume.yaml"
        request = {
            "op": "extract_resume",
            "params": {"path": str(fixture), "fileName": "中文简历.yaml"},
        }
        env = os.environ.copy()
        env["PYTHONUTF8"] = "0"
        env.pop("PYTHONIOENCODING", None)
        process = subprocess.run(
            [sys.executable, str(ROOT / "worker.py")],
            input=(json.dumps(request, ensure_ascii=False) + "\n").encode("utf-8"),
            capture_output=True,
            env=env,
            timeout=15,
            check=False,
        )
        self.assertEqual(process.returncode, 0, process.stderr.decode("utf-8"))
        messages = [json.loads(line) for line in process.stdout.decode("utf-8").splitlines()]
        result = next(message for message in messages if message.get("type") == "result")
        self.assertEqual(result["data"]["profile"]["name"], "林知远")
        self.assertEqual(result["data"]["profile"]["location"], "上海")
        self.assertEqual(result["data"]["profile"]["sourceFileName"], "中文简历.yaml")


if __name__ == "__main__":
    unittest.main()
