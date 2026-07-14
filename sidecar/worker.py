#!/usr/bin/env python3
"""One-line JSON RPC worker used by the Tauri backend.

The process reads one request from stdin and writes JSON Lines to stdout. Vendor
and RenderCV output is captured so stdout always remains machine-readable.
"""

from __future__ import annotations

import contextlib
import atexit
import io
import json
import os
import pathlib
import re
import shutil
import signal
import subprocess
import sys
import tempfile
import traceback
import uuid
from collections import Counter
from datetime import datetime, timedelta, timezone
from typing import Any


SHANGHAI_TZ = timezone(timedelta(hours=8), name="Asia/Shanghai")
APP_VERSION = "0.2.0"
PROTOCOL_VERSION = "2"


def configure_utf8_stdio() -> None:
    """Keep the JSONL boundary UTF-8 even on Chinese Windows code pages."""
    for stream in (sys.stdin, sys.stdout, sys.stderr):
        reconfigure = getattr(stream, "reconfigure", None)
        if callable(reconfigure):
            reconfigure(encoding="utf-8", errors="strict")


configure_utf8_stdio()
PROTOCOL_STDOUT = sys.stdout
_active_boss: Any | None = None
_cleaning_boss = False


def emit(kind: str, **payload: Any) -> bool:
    """Write one protocol event without crashing if the parent closed its pipe."""
    try:
        print(
            json.dumps({"type": kind, **payload}, ensure_ascii=False),
            file=PROTOCOL_STDOUT,
            flush=True,
        )
        return True
    except (OSError, ValueError, AttributeError):
        return False


def print_current_exception() -> None:
    """Best-effort traceback output for windowed/PyInstaller sidecars."""
    try:
        traceback.print_exc(file=sys.stderr)
    except (OSError, ValueError, AttributeError):
        pass


def now() -> str:
    return datetime.now(SHANGHAI_TZ).isoformat()


def split_pipe(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    return [part.strip() for part in str(value or "").split("|") if part.strip()]


def stable_job_id(external_id: str, company: str, title: str, location: str) -> str:
    source = external_id or f"{company}|{title}|{location}"
    return str(uuid.uuid5(uuid.NAMESPACE_URL, f"boss:{source}"))


def normalize_job(raw: dict[str, Any], detail: dict[str, Any] | None = None) -> dict[str, Any]:
    detail = detail or {}
    labels = split_pipe(raw.get("job_labels") or raw.get("tags") or detail.get("tags_list"))
    experience = next((item for item in labels if "年" in item or item in {"应届生", "经验不限", "在校生"}), "")
    degree_values = {"初中", "中专", "高中", "大专", "本科", "硕士", "博士", "学历不限"}
    degree = next((item for item in labels if any(value in item for value in degree_values)), "")
    external_id = str(raw.get("job_id") or raw.get("encrypt_job_id") or "")
    company = str(raw.get("boss_name") or detail.get("company") or "")
    # BOSS list data uses boss_name for the company display in the reference scraper.
    company = str(detail.get("company") or company)
    title = str(raw.get("title") or detail.get("title") or "")
    location = str(raw.get("location") or detail.get("location") or "")
    skills = split_pipe(raw.get("skills"))
    skills.extend(str(item) for item in detail.get("skill_tags", []) if str(item).strip())
    skills = list(dict.fromkeys(skills))
    seen = now()
    return {
        "id": stable_job_id(external_id, company, title, location),
        "source": "boss",
        "externalId": external_id,
        "title": title,
        "company": company,
        "salary": str(raw.get("salary") or detail.get("salary") or "薪资面议"),
        "location": location,
        "experience": experience,
        "degree": degree,
        "companyScale": str(raw.get("company_scale") or ""),
        "companyStage": str(raw.get("company_stage") or ""),
        "industry": str(raw.get("company_industry") or ""),
        "skills": skills,
        "welfare": split_pipe(raw.get("welfare")),
        "description": str(detail.get("jd") or ""),
        "sourceUrl": str(raw.get("job_link") or detail.get("job_link") or detail.get("link") or ""),
        "bossName": None,
        "bossTitle": str(raw.get("boss_title") or "") or None,
        "firstSeen": seen,
        "lastSeen": seen,
        "isNew": True,
        "fit": None,
        "greeting": None,
        "patches": [],
    }


def market_report(jobs: list[dict[str, Any]], keyword: str, city: str) -> str:
    skills = Counter(skill for job in jobs for skill in job.get("skills", []))
    experience = Counter(job.get("experience") for job in jobs if job.get("experience"))
    degrees = Counter(job.get("degree") for job in jobs if job.get("degree"))
    top_skills = "、".join(item for item, _ in skills.most_common(6)) or "数据不足"
    top_experience = experience.most_common(1)[0][0] if experience else "数据不足"
    top_degree = degrees.most_common(1)[0][0] if degrees else "数据不足"
    return (
        f"## 本次岗位观察\n\n"
        f"- 本次整理 **{len(jobs)}** 个“{keyword}”岗位，范围为 {city}。\n"
        f"- 高频技能为 **{top_skills}**。\n"
        f"- 最常见经验要求是 **{top_experience}**，学历要求以 **{top_degree}** 为主。\n\n"
        f"> 建议先在简历前半页突出与高频技能直接相关、可量化的项目成果。"
    )


def load_boss_module():
    try:
        from vendor import boss_cdp_raw as boss
    except ImportError:
        from sidecar.vendor import boss_cdp_raw as boss
    if not boss.require_runtime_dependencies("requests", "websocket"):
        raise RuntimeError("BOSS 抓取依赖缺失，请重新构建 sidecar。")
    return boss


def setup_boss(params: dict[str, Any]) -> dict[str, Any]:
    global _active_boss
    boss = load_boss_module()
    _active_boss = boss
    reset_requested = bool(params.get("resetProfile", False))
    outcome: dict[str, Any] = {
        "loginSucceeded": False,
        "resetRequested": reset_requested,
        "cleanupSucceeded": True,
        "closedProcesses": 0,
        "error": None,
    }

    def record_cleanup(cleanup: dict[str, Any]) -> None:
        outcome["closedProcesses"] += int(cleanup["closedProcesses"])
        if not cleanup["cleanupSucceeded"]:
            outcome["cleanupSucceeded"] = False
            cleanup_error = str(cleanup.get("error") or "无法确认专用 Chrome 已关闭。")
            if not outcome["error"]:
                outcome["error"] = cleanup_error

    try:
        if reset_requested:
            # Never let prepare_cdp_profile delete the profile while one of its
            # Chrome processes is still alive. The check is scoped to the exact
            # --user-data-dir used by the sidecar, not the user's normal Chrome.
            reset_cleanup = close_boss_session(boss, progress=12, action="重新配置前")
            record_cleanup(reset_cleanup)
            if not reset_cleanup["cleanupSucceeded"]:
                raise RuntimeError(
                    f"无法安全重置 BOSS 专用浏览器配置：{reset_cleanup['error']}"
                )

        ensure_boss_session(
            boss,
            int(params.get("loginTimeout", 300)),
            reset_profile=reset_requested,
        )
        outcome["loginSucceeded"] = True
    except Exception as error:  # Setup failures are part of the structured outcome.
        if not outcome["error"]:
            outcome["error"] = str(error)
    finally:
        record_cleanup(close_boss_session(boss, progress=90, action="配置结束后"))
        _active_boss = None

    return outcome


def ensure_boss_session(boss: Any, login_timeout: int, reset_profile: bool = False) -> None:
    with contextlib.redirect_stdout(io.StringIO()) as captured:
        code = boss.run_setup_chrome(
            9222,
            copy_login_state=False,
            reset_profile=reset_profile,
            wait_login=True,
            login_timeout=login_timeout,
        )
    if code != 0:
        raise RuntimeError(captured.getvalue().strip() or "BOSS 登录未完成。")


def close_boss_session(
    boss: Any,
    *,
    progress: int = 90,
    action: str = "任务结束后",
) -> dict[str, Any]:
    """Close and verify only Chrome processes using the isolated BOSS profile."""
    result: dict[str, Any] = {
        "cleanupSucceeded": False,
        "closedProcesses": 0,
        "error": None,
    }
    try:
        stopped = boss.stop_cdp_chrome(boss.DEFAULT_CDP_DATA_DIR)
        result["closedProcesses"] = int(stopped or 0)
        remaining = list(boss.chrome_pids_for_user_data_dir(boss.DEFAULT_CDP_DATA_DIR))
        if remaining:
            raise RuntimeError(f"仍检测到专用 Chrome 进程：{remaining}")
        result["cleanupSucceeded"] = True
        message = f"{action}已关闭并确认 BOSS 专用 Chrome（{stopped} 个进程）"
    except Exception as error:  # Cleanup failure must not discard scraped jobs.
        result["error"] = str(error)
        message = f"{action}关闭 BOSS 专用 Chrome 失败：{error}"
    emit("progress", progress=progress, message=message)
    return result


def close_boss(_params: dict[str, Any]) -> dict[str, Any]:
    return close_boss_session(load_boss_module(), action="手动清理时")


def environment_status(_params: dict[str, Any]) -> dict[str, Any]:
    boss = load_boss_module()
    executable = pathlib.Path(str(boss.DEFAULT_CHROME_PATH)).expanduser()
    installed = executable.is_file()
    version = installed_chrome_version(executable) if installed else None
    return {
        "appVersion": APP_VERSION,
        "protocolVersion": PROTOCOL_VERSION,
        "chrome": {
            "installed": installed,
            "version": version,
            "executablePath": str(executable) if installed else None,
        },
    }


def installed_chrome_version(executable: pathlib.Path) -> str | None:
    """Read the installed Chrome version using platform-appropriate metadata.

    Chrome's Windows ``--version`` output follows the active system code page,
    which made UTF-8 sidecars display mojibake and could also activate a browser
    process. Official Windows installs keep a version-named directory beside
    chrome.exe, so use that static metadata there. macOS and Linux do not use
    the Windows directory layout, and their ``--version`` output is UTF-8 safe.
    """
    if os.name != "nt":
        try:
            completed = subprocess.run(
                [str(executable), "--version"],
                capture_output=True,
                text=True,
                timeout=5,
                check=False,
                encoding="utf-8",
                errors="replace",
            )
        except (OSError, subprocess.SubprocessError):
            return None
        return (completed.stdout or completed.stderr).strip() or None

    version_pattern = re.compile(r"^\d+(?:\.\d+){3}$")
    try:
        versions = [
            child.name
            for child in executable.parent.iterdir()
            if child.is_dir() and version_pattern.fullmatch(child.name)
        ]
    except OSError:
        return None
    if not versions:
        return None
    return max(versions, key=lambda value: tuple(int(part) for part in value.split(".")))


def clear_boss_data(_params: dict[str, Any]) -> dict[str, Any]:
    boss = load_boss_module()
    outcome = close_boss_session(boss, action="清除数据前")
    if not outcome["cleanupSucceeded"]:
        raise RuntimeError(outcome.get("error") or "无法确认 BOSS 专用 Chrome 已关闭")
    remaining = list(boss.chrome_pids_for_user_data_dir(boss.DEFAULT_CDP_DATA_DIR))
    if remaining:
        raise RuntimeError(f"仍检测到 BOSS 专用 Chrome 进程：{remaining}")
    deleted: list[str] = []
    for raw_path in (boss.DEFAULT_CDP_DATA_DIR, boss.DEFAULT_RESULT_DIR):
        path = pathlib.Path(str(raw_path)).expanduser()
        if path.exists():
            shutil.rmtree(path)
            deleted.append(str(path))
    return {"deleted": deleted, "remainingPids": []}


def scrape_jobs(params: dict[str, Any]) -> dict[str, Any]:
    global _active_boss
    if not str(params.get("keyword") or "").strip():
        raise ValueError("岗位关键词不能为空。")
    boss = load_boss_module()
    _active_boss = boss
    try:
        return _scrape_jobs(boss, params)
    finally:
        # Login timeouts, captcha errors, list/detail exceptions, and empty
        # results all travel through this path. Cleanup errors are reported as
        # progress but do not hide the original scrape result or exception.
        close_boss_session(boss)
        _active_boss = None


def _scrape_jobs(boss: Any, params: dict[str, Any]) -> dict[str, Any]:
    keyword = str(params.get("keyword") or "").strip()
    if not keyword:
        raise ValueError("岗位关键词不能为空。")
    city = str(params.get("city") or "上海").strip()
    pages = max(1, min(int(params.get("pages") or 1), 5))
    completed_detail_external_ids = {
        str(external_id).strip()
        for external_id in (params.get("completedDetailExternalIds") or [])
        if str(external_id).strip()
    }
    city_name, city_code = boss.resolve_city(city)
    city_code = str(city_code).strip()
    if not re.fullmatch(r"\d{9}", city_code):
        raise RuntimeError(f"无法识别城市“{city}”，请填写城市中文名（例如：上海）或 9 位 BOSS 城市代码。")
    filters: dict[str, str] = {}
    mapping = {"salary": "salary", "experience": "experience", "degree": "degree", "companyScale": "scale"}
    for input_name, scraper_name in mapping.items():
        if params.get(input_name):
            filters[scraper_name] = str(params[input_name])
    emit("progress", progress=12, message="正在检查 BOSS 登录状态；若出现登录界面，请完成登录")
    ensure_boss_session(boss, int(params.get("loginTimeout", 300)))
    emit("progress", progress=24, message=f"登录状态正常，城市：{city_name}（{city_code}）")
    with tempfile.TemporaryDirectory(prefix="ai-job-scrape-") as temporary:
        output = str(pathlib.Path(temporary) / "jobs.json")
        details_output = str(pathlib.Path(temporary) / "details.json")
        with contextlib.redirect_stdout(io.StringIO()):
            listing = boss.scrape_list(
                keyword,
                city_code,
                pages,
                filters,
                output,
                cdp_port=9222,
                fmt="json",
                allow_dom_fallback=False,
                on_job=lambda raw: emit("job", phase="list", job=normalize_job(raw)),
            )

            listed_jobs = list(listing.get("jobs") or [])
            detail_candidates = [
                raw
                for raw in listed_jobs
                if str(raw.get("job_id") or raw.get("encrypt_job_id") or "").strip()
                not in completed_detail_external_ids
            ]
            skipped_existing = len(listed_jobs) - len(detail_candidates)
            detail_state = {
                "succeeded": 0,
                "skipped": skipped_existing,
                "failed": 0,
                "processed": skipped_existing,
            }

            emit(
                "progress",
                progress=55,
                message=(
                    f"岗位列表抓取完成，共 {len(listed_jobs)} 个；"
                    f"已有详情跳过 {skipped_existing} 个"
                ),
                detailTotal=len(listed_jobs),
                detailProcessed=skipped_existing,
                detailSucceeded=0,
                detailSkipped=skipped_existing,
                detailFailed=0,
            )

            def emit_detail(raw: dict[str, Any], detail: dict[str, Any]) -> None:
                if str(detail.get("jd") or "").strip():
                    emit("job", phase="detail", job=normalize_job(raw, detail))

            def emit_detail_progress(**detail_progress: Any) -> None:
                detail_state["succeeded"] = int(detail_progress.get("succeeded") or 0)
                detail_state["skipped"] = skipped_existing + int(detail_progress.get("skipped") or 0)
                detail_state["failed"] = int(detail_progress.get("failed") or 0)
                detail_state["processed"] = skipped_existing + int(detail_progress.get("processed") or 0)
                total = len(listed_jobs)
                progress = 55 + int(22 * detail_state["processed"] / max(total, 1))
                emit(
                    "progress",
                    progress=min(progress, 77),
                    message=(
                        "正在抓取岗位详情："
                        f"成功 {detail_state['succeeded']}，"
                        f"跳过 {detail_state['skipped']}，"
                        f"失败 {detail_state['failed']}"
                        f"（{detail_state['processed']}/{total}）"
                    ),
                    detailTotal=total,
                    detailProcessed=detail_state["processed"],
                    detailSucceeded=detail_state["succeeded"],
                    detailSkipped=detail_state["skipped"],
                    detailFailed=detail_state["failed"],
                )

            detail_listing = {**listing, "jobs": detail_candidates, "total": len(detail_candidates)}
            details = boss.scrape_details(
                detail_listing,
                None,
                details_output,
                cdp_port=9222,
                fmt="json",
                on_detail=emit_detail,
                on_progress=emit_detail_progress,
            ) if detail_candidates else []

    if not listing.get("jobs"):
        raise RuntimeError(
            f"BOSS 未返回岗位：关键词“{keyword}”，城市“{city_name}”（{city_code}）。"
            "请确认登录未过期、页面没有验证码，并尝试更宽泛的关键词。"
        )

    detail_by_id = {str(detail.get("job_id")): detail for detail in details or []}
    jobs = [normalize_job(raw, detail_by_id.get(str(raw.get("job_id")))) for raw in listing.get("jobs", [])]
    emit(
        "progress",
        progress=78,
        message=(
            "岗位详情抓取完成："
            f"成功 {detail_state['succeeded']}，"
            f"跳过 {detail_state['skipped']}，"
            f"失败 {detail_state['failed']}"
        ),
        detailTotal=len(listing.get("jobs") or []),
        detailProcessed=detail_state["processed"],
        detailSucceeded=detail_state["succeeded"],
        detailSkipped=detail_state["skipped"],
        detailFailed=detail_state["failed"],
    )
    return {
        "jobs": jobs,
        "reportMarkdown": market_report(jobs, keyword, city_name),
        "resolvedCity": city_name,
        "cityCode": city_code,
        "detailSummary": detail_state,
    }


def extract_docx_text(path: pathlib.Path) -> str:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = Document(str(path))
    lines: list[str] = []
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            value = Paragraph(child, document).text.strip()
            if value:
                lines.append(value)
        elif isinstance(child, CT_Tbl):
            table = Table(child, document)
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                value = " | ".join(dict.fromkeys(cell for cell in cells if cell))
                if value:
                    lines.append(value)
    for section in document.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                value = paragraph.text.strip()
                if value and value not in lines:
                    lines.append(value)
            for table in container.tables:
                for row in table.rows:
                    value = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip())
                    if value and value not in lines:
                        lines.append(value)
    return "\n".join(lines)


def render_pdf_page(path: pathlib.Path, page_index: int, output_dir: pathlib.Path) -> pathlib.Path:
    import pypdfium2 as pdfium

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"page-{page_index + 1}.png"
    with contextlib.closing(pdfium.PdfDocument(str(path))) as document:
        with contextlib.closing(document[page_index]) as page:
            with contextlib.closing(page.render(scale=2.25)) as bitmap:
                with contextlib.closing(bitmap.to_pil()) as image:
                    image.save(output_path, format="PNG", optimize=True)
    return output_path


def extract_text(path: pathlib.Path) -> tuple[str, dict[str, Any] | None, list[dict[str, Any]]]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        from pypdf import PdfReader
        pages: list[dict[str, Any]] = []
        page_texts: list[str] = []
        reader = PdfReader(str(path))
        image_dir = path.parent / f"{path.stem}-pages"
        for index, page in enumerate(reader.pages):
            page_text = (page.extract_text() or "").strip()
            page_texts.append(page_text)
            page_data: dict[str, Any] = {"pageNumber": index + 1, "text": page_text}
            has_visual_content = bool(list(page.images))
            if len(page_text) < 30 and has_visual_content:
                page_data["imagePath"] = str(render_pdf_page(path, index, image_dir).resolve())
            pages.append(page_data)
        return "\n\n".join(page_texts), None, pages
    if suffix == ".docx":
        text = extract_docx_text(path)
        if len(text.strip()) < 20:
            raise RuntimeError("DOCX 中没有足够的可读取文本。")
        return text, None, []
    if suffix in {".yaml", ".yml"}:
        import yaml
        data = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
        return path.read_text(encoding="utf-8"), data, []
    raise RuntimeError("仅支持 PDF、DOCX、YAML 和 YML。")


def first_value(data: dict[str, Any], *keys: str) -> str:
    for key in keys:
        value = data.get(key)
        if value is not None:
            return str(value)
    return ""


def split_skill_items(value: Any) -> list[str]:
    if isinstance(value, list):
        values = [str(item).strip() for item in value]
    else:
        values = [item.strip() for item in re.split(r"[,，、|]", str(value or ""))]
    return list(dict.fromkeys(item for item in values if item))


def entry_dates(entry: dict[str, Any]) -> tuple[str, str]:
    start = first_value(entry, "start_date", "startDate")
    end = first_value(entry, "end_date", "endDate")
    combined = first_value(entry, "date")
    if combined and not (start or end):
        start, end = normalize_date_pair("", combined)
    return normalize_date_pair(start, end)


DATE_RANGE_PATTERN = re.compile(
    r"^\s*(\d{4}(?:[./\-年]\d{1,2}(?:月)?)?)\s*(?:-|–|—|至|到)\s*(\d{4}(?:[./\-年]\d{1,2}(?:月)?)?|至今|现在|present)\s*$",
    re.IGNORECASE,
)


def normalize_date_pair(start: Any, end: Any) -> tuple[str, str]:
    start_value = str(start or "").strip().strip("-–— ")
    end_value = str(end or "").strip().strip("-–— ")
    candidate = end_value if not start_value else start_value if not end_value else ""
    match = DATE_RANGE_PATTERN.fullmatch(candidate) if candidate else None
    if match:
        return match.group(1).strip(), match.group(2).strip()
    return start_value, end_value


def render_date_fields(item: dict[str, Any]) -> tuple[str | None, str | None]:
    start, end = normalize_date_pair(item.get("startDate"), item.get("endDate"))
    if not start and end:
        return rendercv_date(end), None
    return rendercv_date(start), rendercv_date(end, end_date=True)


def rendercv_date(value: str, *, end_date: bool = False) -> str | None:
    normalized = value.strip()
    if not normalized:
        return None
    if end_date and normalized.lower() in {"至今", "现在", "present", "current"}:
        return "present"
    match = re.fullmatch(r"(\d{4})(?:[./年-](\d{1,2})(?:月)?)?(?:[./日-](\d{1,2})(?:日)?)?", normalized)
    if not match:
        return normalized
    return "-".join(part.zfill(2) if index else part for index, part in enumerate(match.groups(default="")) if part)


def rendercv_phone(value: Any) -> str | None:
    phone = str(value or "").strip()
    digits = re.sub(r"\D", "", phone)
    if len(digits) == 11 and digits.startswith("1"):
        return f"+86{digits}"
    return phone or None


def display_degree(item: dict[str, Any]) -> str:
    degree = str(item.get("degree") or "").strip()
    if degree == "其他":
        return str(item.get("degreeDetail") or "").strip() or degree
    return degree


def profile_from_yaml(data: dict[str, Any], file_name: str) -> dict[str, Any]:
    cv = data.get("cv") if isinstance(data.get("cv"), dict) else data
    sections = cv.get("sections") if isinstance(cv.get("sections"), dict) else {}
    experience_entries: list[dict[str, Any]] = []
    education_entries: list[dict[str, Any]] = []
    professional_skills: list[dict[str, Any]] = []
    projects: list[dict[str, Any]] = []
    certifications: list[dict[str, Any]] = []
    summary = ""
    for section_name, entries in sections.items():
        name = str(section_name).lower()
        entries = entries if isinstance(entries, list) else []
        if any(word in name for word in ["education", "教育"]):
            for entry in entries:
                if isinstance(entry, dict):
                    start_date, end_date = entry_dates(entry)
                    education_entries.append({
                        "institution": first_value(entry, "institution"), "area": first_value(entry, "area"),
                        "degree": first_value(entry, "degree"), "startDate": start_date,
                        "degreeDetail": first_value(entry, "degree_detail", "degreeDetail"),
                        "endDate": end_date, "highlights": [str(item) for item in entry.get("highlights", [])],
                    })
        elif any(word in name for word in ["experience", "工作", "职业经历"]) and not any(word in name for word in ["project", "项目"]):
            for entry in entries:
                if isinstance(entry, dict):
                    start_date, end_date = entry_dates(entry)
                    experience_entries.append({
                        "company": first_value(entry, "company", "institution"),
                        "position": first_value(entry, "position", "title"),
                        "location": first_value(entry, "location"),
                        "startDate": start_date,
                        "endDate": end_date,
                        "highlights": [str(item) for item in entry.get("highlights", [])],
                    })
        elif "skill" in name or "技能" in name:
            for entry in entries:
                if isinstance(entry, str):
                    professional_skills.append({"id": str(uuid.uuid4()), "label": "核心技能", "items": split_skill_items(entry)})
                elif isinstance(entry, dict):
                    professional_skills.append({
                        "id": str(uuid.uuid4()),
                        "label": first_value(entry, "label") or "专业技能",
                        "items": split_skill_items(entry.get("details") or entry.get("items")),
                    })
        elif any(word in name for word in ["project", "项目"]):
            for entry in entries:
                if isinstance(entry, dict):
                    start_date, end_date = entry_dates(entry)
                    projects.append({
                        "id": str(uuid.uuid4()), "name": first_value(entry, "name", "title"),
                        "summary": first_value(entry, "summary"), "startDate": start_date, "endDate": end_date,
                        "highlights": [str(item) for item in entry.get("highlights", [])],
                    })
        elif any(word in name for word in ["certification", "certificate", "证书", "资质"]):
            for entry in entries:
                if isinstance(entry, dict):
                    certifications.append({
                        "id": str(uuid.uuid4()), "name": first_value(entry, "name", "title"),
                        "issuer": first_value(entry, "issuer", "institution"), "date": first_value(entry, "date"),
                    })
                elif str(entry).strip():
                    certifications.append({"id": str(uuid.uuid4()), "name": str(entry).strip(), "issuer": "", "date": ""})
        elif any(word in name for word in ["summary", "profile", "简介", "个人定位"]):
            summary = " ".join(str(item) for item in entries)
    return base_profile(file_name, first_value(cv, "name"), first_value(cv, "email"), first_value(cv, "phone"), first_value(cv, "location"), first_value(cv, "website"), first_value(cv, "headline"), summary, professional_skills, experience_entries, education_entries, projects, certifications)


KNOWN_SKILLS = ["Python", "Java", "Golang", "Rust", "TypeScript", "JavaScript", "Svelte", "React", "Vue", "FastAPI", "Django", "Flask", "LangChain", "RAG", "Docker", "Kubernetes", "Redis", "MySQL", "PostgreSQL", "PyTorch", "TensorFlow", "AWS", "Azure"]


def profile_from_text(text: str, file_name: str) -> dict[str, Any]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    email_match = re.search(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", text)
    phone_match = re.search(r"(?<!\d)(?:\+?86[- ]?)?1[3-9]\d(?:[- ]?\d){8}(?!\d)", text)
    name = next((line for line in lines[:8] if 1 < len(line) <= 12 and not re.search(r"[@\d:/]", line)), "")
    skills = [skill for skill in KNOWN_SKILLS if re.search(rf"(?<![A-Za-z]){re.escape(skill)}(?![A-Za-z])", text, re.IGNORECASE)]
    headline = next((line for line in lines[:12] if any(word in line.lower() for word in ["工程师", "开发", "产品", "designer", "engineer"])), "")
    summary = next((line for line in lines if 35 <= len(line) <= 180), "")
    skill_groups = [{"id": str(uuid.uuid4()), "label": "核心技能", "items": skills}] if skills else []
    return base_profile(file_name, name, email_match.group(0) if email_match else "", phone_match.group(0) if phone_match else "", "", "", headline, summary, skill_groups, [], [])


def base_profile(file_name: str, name: str, email: str, phone: str, location: str, website: str, headline: str, summary: str, professional_skills: list[dict[str, Any]], experiences: list[dict[str, Any]], education: list[dict[str, Any]], projects: list[dict[str, Any]] | None = None, certifications: list[dict[str, Any]] | None = None) -> dict[str, Any]:
    if professional_skills and isinstance(professional_skills[0], str):
        professional_skills = [{"id": str(uuid.uuid4()), "label": "核心技能", "items": split_skill_items(professional_skills)}]
    professional_skills = [group for group in professional_skills if group.get("items")]
    facts = [
        {"id": str(uuid.uuid4()), "category": "skill", "value": skill, "source": f"{file_name} · 专业技能 · {group.get('label', '')}", "confidence": 0.95, "confirmed": False}
        for group in professional_skills for skill in group.get("items", [])
    ]
    return {
        "id": "resume-master", "name": name, "headline": headline, "email": email, "phone": phone,
        "location": location, "website": website, "summary": summary, "templateId": "ai-engineering",
        "professionalSkills": professional_skills, "experiences": experiences, "education": education,
        "projects": projects or [], "certifications": certifications or [], "facts": facts,
        "preferences": {"targetRoles": [], "cities": [], "remotePreference": "flexible", "energizingTasks": [], "drainingTasks": [], "hardConstraints": []},
        "sourceFileName": file_name, "updatedAt": now(), "version": 1,
    }


def extract_resume(params: dict[str, Any]) -> dict[str, Any]:
    path = pathlib.Path(str(params["path"]))
    file_name = str(params.get("fileName") or path.name)
    text, yaml_data, pages = extract_text(path)
    profile = profile_from_yaml(yaml_data, file_name) if isinstance(yaml_data, dict) else profile_from_text(text, file_name)
    return {"profile": profile, "rawText": text, "pages": pages}


RESUME_COLOR_THEMES = {
    "pine": {"accent": "#176B57", "links": "#0B7A67"},
    "navy": {"accent": "#1F407A", "links": "#005CB8"},
    "graphite": {"accent": "#24292F", "links": "#24292F"},
}

RESUME_BOLD_KEYWORDS = [
    "Dify", "FastAPI", "Docker", "Docker Compose", "PostgreSQL", "vLLM", "SGLang",
    "llama.cpp", "MinerU", "Milvus", "OpenAI", "Linux", "Prometheus", "Grafana",
    "Triton", "PP-OCRv6", "PP-StructureV3", "PaddleOCR-VL-1.6",
]


def resume_design(color_theme: str) -> dict[str, Any]:
    colors = RESUME_COLOR_THEMES.get(color_theme)
    if colors is None:
        raise ValueError("不支持的简历颜色主题。")
    accent = colors["accent"]
    return {
        "theme": "classic",
        "page": {
            "size": "a4",
            "top_margin": "1.2cm",
            "bottom_margin": "1.2cm",
            "left_margin": "1.35cm",
            "right_margin": "1.35cm",
            "show_footer": False,
            "show_top_note": False,
        },
        "colors": {
            "name": accent,
            "headline": accent,
            "connections": accent,
            "section_titles": accent,
            "links": colors["links"],
        },
        "typography": {
            "line_spacing": "0.72em",
            "alignment": "left",
            "font_family": {
                "body": "Microsoft YaHei",
                "name": "Microsoft YaHei",
                "headline": "Microsoft YaHei",
                "connections": "Microsoft YaHei",
                "section_titles": "Microsoft YaHei",
            },
            "font_size": {
                "body": "10.2pt",
                "name": "25pt",
                "headline": "10.6pt",
                "connections": "9.8pt",
                "section_titles": "1.28em",
            },
        },
        "header": {
            "alignment": "center",
            "space_below_name": "0.22cm",
            "space_below_headline": "0.24cm",
            "space_below_connections": "0.32cm",
            "connections": {
                "phone_number_format": "international",
                "show_icons": False,
                "separator": "|",
                "space_between_connections": "0.32cm",
            },
        },
        "section_titles": {
            "type": "with_full_line",
            "space_above": "0.42cm",
            "space_below": "0.22cm",
        },
        "sections": {
            "space_between_regular_entries": "0.42cm",
            "space_between_text_based_entries": "0.14cm",
            "show_time_spans_in": [],
        },
        "entries": {
            "date_and_location_width": "4.6cm",
            "side_space": "0cm",
            "space_between_columns": "0.24cm",
            "allow_page_break": False,
            "short_second_row": False,
            "summary": {"space_above": "0.06cm"},
            "highlights": {
                "space_left": "0.05cm",
                "space_above": "0.08cm",
                "space_between_items": "0.06cm",
                "space_between_bullet_and_text": "0.32em",
            },
        },
        "templates": {
            "experience_entry": {
                "main_column": "**COMPANY**, POSITION\nSUMMARY\nHIGHLIGHTS",
                "date_and_location_column": "LOCATION · DATE",
            },
            "education_entry": {
                "main_column": "**INSTITUTION**, AREA\nSUMMARY\nHIGHLIGHTS",
                "degree_column": "**DEGREE**",
                "date_and_location_column": "LOCATION · DATE",
            },
        },
    }


def profile_to_rendercv(profile: dict[str, Any], color_theme: str = "navy") -> dict[str, Any]:
    section_values: dict[str, tuple[str, Any]] = {}
    if profile.get("summary"):
        section_values["summary"] = ("个人简介", [profile["summary"]])
    skill_groups = profile.get("professionalSkills") or []
    if not skill_groups and profile.get("skills"):
        skill_groups = [{"label": "核心技能", "items": profile["skills"]}]
    if skill_groups:
        section_values["professionalSkills"] = ("专业技能", [
            {"label": group.get("label") or "专业技能", "details": ", ".join(group.get("items") or [])}
            for group in skill_groups if group.get("items")
        ])
    if profile.get("experiences"):
        experience_entries = []
        for item in profile["experiences"]:
            start_date, end_date = render_date_fields(item)
            experience_entries.append({
                "company": item.get("company", ""), "position": item.get("position", ""), "location": item.get("location", ""),
                "start_date": start_date, "end_date": end_date,
                "highlights": item.get("highlights", []),
            })
        section_values["experiences"] = ("工作经历", experience_entries)
    if profile.get("projects"):
        project_entries = []
        for item in profile["projects"]:
            start_date, end_date = render_date_fields(item)
            project_entries.append({
                "name": item.get("name", ""), "summary": item.get("summary", ""),
                "start_date": start_date, "end_date": end_date,
                "highlights": item.get("highlights", []),
            })
        section_values["projects"] = ("项目经历", project_entries)
    if profile.get("certifications"):
        section_values["certifications"] = ("证书 / 专业资质", [
            " · ".join(part for part in [item.get("name", ""), item.get("issuer", ""), item.get("date", "")] if part)
            for item in profile["certifications"]
        ])
    if profile.get("education"):
        education_entries = []
        for item in profile["education"]:
            start_date, end_date = render_date_fields(item)
            education_entries.append({
                "institution": item.get("institution", ""), "area": item.get("area", ""), "degree": display_degree(item),
                "start_date": start_date, "end_date": end_date,
                "highlights": item.get("highlights", []),
            })
        section_values["education"] = ("教育经历", education_entries)
    orders = {
        "ai-engineering": ["summary", "professionalSkills", "projects", "experiences", "certifications", "education"],
        "data-analysis": ["summary", "professionalSkills", "experiences", "projects", "certifications", "education"],
        "finance-accounting": ["summary", "experiences", "certifications", "professionalSkills", "education", "projects"],
        "general": ["summary", "experiences", "professionalSkills", "projects", "certifications", "education"],
    }
    sections: dict[str, Any] = {}
    for key in orders.get(str(profile.get("templateId") or "ai-engineering"), orders["ai-engineering"]):
        if key in section_values and section_values[key][1]:
            title, entries = section_values[key]
            sections[title] = entries
    return {
        "cv": {
            "name": profile.get("name") or "Candidate", "headline": profile.get("headline") or None,
            "location": profile.get("location") or None, "email": profile.get("email") or None,
            "phone": rendercv_phone(profile.get("phone")), "website": profile.get("website") or None,
            "sections": sections,
        },
        "design": resume_design(color_theme),
        "locale": {"language": "mandarin_chinese"},
        "settings": {"bold_keywords": RESUME_BOLD_KEYWORDS},
    }


def render_resume(params: dict[str, Any]) -> dict[str, Any]:
    try:
        import yaml
        from rendercv.renderer import pdf_png
        from rendercv.renderer.typst import generate_typst
        from rendercv.schema.rendercv_model_builder import build_rendercv_dictionary_and_model
    except ImportError as error:
        raise RuntimeError("RenderCV 运行时未安装，请使用生产 sidecar 或安装 sidecar 依赖。") from error

    output_path = pathlib.Path(str(params["outputPath"])).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    data = profile_to_rendercv(dict(params["profile"]), str(params.get("colorTheme") or "navy"))
    yaml_text = yaml.safe_dump(data, allow_unicode=True, sort_keys=False)
    with tempfile.TemporaryDirectory(prefix="resume-render-") as temporary:
        temporary_path = pathlib.Path(temporary)
        yaml_path = temporary_path / "resume.yaml"
        typst_path = temporary_path / "resume.typ"
        yaml_path.write_text(yaml_text, encoding="utf-8")
        _, model = build_rendercv_dictionary_and_model(
            yaml_text,
            input_file_path=yaml_path,
            output_folder=temporary_path,
            typst_path=typst_path,
            pdf_path=output_path,
            dont_generate_png=True,
            dont_generate_markdown=True,
            dont_generate_html=True,
        )
        # RenderCV 2.8 imports Font Awesome even when a theme disables icons.
        # get_package_path() returns a process-private temporary copy. Patch that
        # copy atomically so the installed RenderCV package is never modified.
        package_path = pdf_png.get_package_path()
        for library in package_path.glob("preview/rendercv/*/lib.typ"):
            source = library.read_text(encoding="utf-8")
            source = source.replace(
                '#import "@preview/fontawesome:0.6.0": fa-icon',
                '#let fa-icon(name, size: 1em) = none',
            )
            replacement = library.with_suffix(".typ.tmp")
            replacement.write_text(source, encoding="utf-8")
            os.replace(replacement, library)
        pdf_png.get_typst_compiler.cache_clear()

        generated_typst = generate_typst(model)
        generated_pdf = pdf_png.generate_pdf(model, generated_typst)
        if generated_pdf is None or not pathlib.Path(generated_pdf).exists():
            raise RuntimeError("RenderCV 没有生成 PDF，请检查简历字段。")
    return {"path": str(output_path)}


OPERATIONS = {
    "setup_boss": setup_boss,
    "close_boss": close_boss,
    "clear_boss_data": clear_boss_data,
    "environment_status": environment_status,
    "scrape_jobs": scrape_jobs,
    "extract_resume": extract_resume,
    "render_resume": render_resume,
    "ping": lambda params: {"python": sys.version, "ok": True},
}


def cleanup_active_boss() -> None:
    global _active_boss, _cleaning_boss
    if _active_boss is None or _cleaning_boss:
        return
    _cleaning_boss = True
    try:
        close_boss_session(_active_boss, action="进程退出前")
    except Exception:  # noqa: BLE001 - shutdown cleanup is best effort
        print_current_exception()
    finally:
        _active_boss = None
        _cleaning_boss = False


def handle_sigterm(_signum: int, _frame: Any) -> None:
    cleanup_active_boss()
    raise SystemExit(143)


atexit.register(cleanup_active_boss)
if hasattr(signal, "SIGTERM"):
    signal.signal(signal.SIGTERM, handle_sigterm)


def main() -> int:
    line = sys.stdin.readline()
    if not line:
        emit("result", ok=False, error="没有收到请求。")
        return 2
    try:
        request = json.loads(line)
        operation = str(request.get("op") or "")
        handler = OPERATIONS.get(operation)
        if handler is None:
            raise RuntimeError(f"未知 sidecar 操作：{operation}")
        emit("progress", progress=5, message=f"开始 {operation}")
        data = handler(dict(request.get("params") or {}))
        emit("result", ok=True, data=data)
        return 0
    except Exception as error:  # noqa: BLE001 - process boundary must serialize every failure
        print_current_exception()
        emit("result", ok=False, error=str(error))
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
